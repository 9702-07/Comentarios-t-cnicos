#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Lógica de extracción y generación para Comentarios Técnicos – Pacific Control SAC.
Usado por la interfaz web (app.py).
"""

import os, re, shutil
from datetime import datetime
import pdfplumber
import openpyxl
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Rutas a archivos fijos (mismo directorio que este archivo)
_HERE            = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH    = os.path.join(_HERE, 'Comentarios técnico ejemplo.docx')
TEMPLATE_SF_PATH = os.path.join(_HERE, 'template_sin_fondo.docx')
LMP_PATH         = os.path.join(_HERE, 'DECRETOS.xlsx')
_FONT_REGULAR    = os.path.join(_HERE, 'Arial.ttf')
_FONT_BOLD       = os.path.join(_HERE, 'ArialBold.ttf')

MESES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
}


# ═══════════════════════════════════════════════════════════════════════════════
# NORMALIZACIÓN
# ═══════════════════════════════════════════════════════════════════════════════

def normalizar(s):
    """Minúsculas, sin acentos, sin paréntesis, espacios colapsados."""
    import unicodedata
    s = str(s).lower().strip()
    s = unicodedata.normalize('NFKD', s)
    s = ''.join(c for c in s if not unicodedata.combining(c))
    s = re.sub(r'\s*\(nmp\)\s*', '', s)
    s = re.sub(r'\(.*?\)', '', s)
    s = re.sub(r'\.\s+', ' ', s)   # "E. coli" → "E coli"
    s = re.sub(r'\.$', '', s)       # punto al final
    s = re.sub(r'\s+', ' ', s).strip()
    return s


# ═══════════════════════════════════════════════════════════════════════════════
# EXTRACCIÓN DE PDFs (Informes de Ensayo)
# ═══════════════════════════════════════════════════════════════════════════════

def extraer_pdf(path):
    """
    Extrae encabezado y tabla de resultados de un Informe de Ensayo de Pacific Control.
    Retorna dict con: numero, razon_social, direccion, procedencia, cotizacion,
                      producto, punto_muestreo, presentacion, tiene_micro, resultados.
    """
    info = {
        'numero': '', 'razon_social': '', 'direccion': '',
        'procedencia': '', 'cotizacion': '', 'producto': '',
        'punto_muestreo': '', 'presentacion': '',
        'tiene_micro': False,
        'resultados': []   # [(analisis, unidad, resultado_str)]
    }

    texto = ''
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            texto += (page.extract_text() or '') + '\n'

    def buscar(pat):
        m = re.search(pat, texto, re.IGNORECASE | re.DOTALL)
        return re.sub(r'\s+', ' ', m.group(1)).strip() if m else ''

    info['numero']         = buscar(r'INFORME DE ENSAYO N[°º]\s*([\d\-]+)')
    info['razon_social']   = buscar(r'Raz[oó]n social del cliente:\s*(.+?)(?:\s*RUC:|\n)')
    info['direccion']      = buscar(r'Domicilio legal del cliente:\s*(.+?)(?:\s*Cotizaci[oó]n:|\n)')
    info['cotizacion']     = buscar(r'Cotizaci[oó]n:\s*([\d\-]+)')
    info['procedencia']    = buscar(r'Procedencia:\s*(.+?)(?:\n|Condici)')
    info['producto']       = buscar(r'Producto declarado:\s*(.+?)(?:\n|N[°º]mero)')
    info['punto_muestreo'] = buscar(r'Punto de muestreo:\s*(.+?)(?:\n|Fecha)')
    info['presentacion']   = buscar(r'Presentaci[oó]n:\s*(.+?)(?:\n|Procedencia)')
    info['tiene_micro']    = bool(re.search(
        r'Microbiol[oó]gic|Coliform|Escherichia|bacterias hetero', texto, re.I))

    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            _extraer_tabla_pagina(page, info['resultados'])

    info['resultados'] = _fusionar_continuaciones(info['resultados'])
    return info


def _extraer_tabla_pagina(page, acum):
    words = page.extract_words(x_tolerance=3, y_tolerance=3, keep_blank_chars=False)
    if not words:
        return

    by_line = {}
    for w in words:
        key = round(w['top'] / 2) * 2
        by_line.setdefault(key, []).append(w)

    col = {}
    in_table = False

    for y in sorted(by_line):
        ws = sorted(by_line[y], key=lambda w: w['x0'])
        line = ' '.join(w['text'] for w in ws)

        # Encabezado de tabla
        if re.search(r'An[aá]lisis', line) and 'Resultados' in line:
            in_table = True
            col = {}
            for w in ws:
                t = w['text']
                if re.match(r'An[aá]lisis', t): col['a']   = w['x0']
                elif t == 'L.D.M':              col['ldm'] = w['x0']
                elif t == 'L.C.M':              col['lcm'] = w['x0']
                elif t == 'Unidad':             col['u']   = w['x0']
                elif t == 'Resultados':         col['r']   = w['x0']
            continue

        if not in_table or not col:
            continue

        if re.search(r'^L\.[DC]\.M\.|^Tipo de an[aá]|^Observac|^FIN|EL USO|P[áa]gina\s+\d', line):
            in_table = False
            col = {}
            continue

        thresh_ldm = col.get('ldm', 180) - 5
        thresh_u   = col.get('u', 273)   - 15
        thresh_r   = col.get('r', 430)   - 60

        nombre_ws, unidad_ws, resultado_ws = [], [], []
        for w in ws:
            x = w['x0']
            if x >= thresh_r:        resultado_ws.append(w['text'])
            elif x >= thresh_u:      unidad_ws.append(w['text'])
            elif x < thresh_ldm:     nombre_ws.append(w['text'])

        nombre    = ' '.join(nombre_ws).strip()
        unidad    = ' '.join(unidad_ws).strip()
        resultado = ' '.join(resultado_ws).strip()

        if not nombre:
            continue
        if nombre in ('Análisis', 'Analisis', 'L.D.M', 'L.C.M'):
            continue
        if re.match(r'^P[áa]gina\s+\d|^FR-|^Our\s|^Offices', nombre):
            continue
        if nombre.startswith('"') or re.search(
                r'INACAL|informe de ensayo.{0,20}al estar|miembros firmantes', nombre, re.I):
            continue

        acum.append((nombre, unidad, resultado))


def _fusionar_continuaciones(filas):
    """Une líneas sin unidad/resultado al nombre de la fila anterior."""
    resultado = []
    for nombre, unidad, valor in filas:
        if unidad == '' and valor == '' and nombre and resultado:
            prev_n, prev_u, prev_v = resultado[-1]
            resultado[-1] = (prev_n + ' ' + nombre, prev_u, prev_v)
        elif valor or nombre:
            resultado.append((nombre, unidad, valor))
    return [(n, u, v) for n, u, v in resultado if v]


# ═══════════════════════════════════════════════════════════════════════════════
# LECTURA DE LMP (Excel o PDF)
# ═══════════════════════════════════════════════════════════════════════════════

def cargar_lmp(path):
    """
    Lee los LMP desde un archivo Excel (.xlsx/.xls) o PDF.
    Retorna dict { normalizar(nombre): (lmp_valor, unidad) }.
    """
    ext = os.path.splitext(path)[1].lower()
    if ext in ('.xlsx', '.xls'):
        return _lmp_desde_excel(path)
    elif ext == '.pdf':
        return _lmp_desde_pdf(path)
    else:
        raise ValueError(f'Formato no soportado para LMP: {ext}. Usa .xlsx o .pdf')


# ── Alias: nombre en DECRETOS Excel → nombres equivalentes en PDFs de Pacific Control ──
# Clave: fragmento normalizado del nombre en el Excel del D.S.
# Valor: lista de nombres normalizados tal como aparecen en los PDFs
DECRETOS_ALIAS = {
    # Bacteriológicos
    'bacterias coliformes totales': [
        'coliformes totales'],
    'e coli':   ['escherichia coli'],
    'e  coli':  ['escherichia coli'],
    'bacterias coliformes termotolerantes': [
        'coliformes fecales o termotolerantes',
        'coliformes fecales o termotolerantes nmp'],
    'bacterias heterotroficas': [
        'recuento de bacterias heterotrofas'],
    # Turbiedad / Turbidez (dos nombres para lo mismo)
    'turbiedad': ['turbidez'],
    'turbidez':  ['turbiedad'],
    # pH
    'ph': [
        'potencial de hidrogeno ph', 'potencial de hidrogeno'],
    # Helmintos y parásitos — agrupados en D.S., individuales en Pacific Control
    'huevos y larvas de helmintos': [
        'trematodos - paragonimus sp.',  'trematodos paragonimus sp',
        'trematodos - fasciola hepatica','trematodos fasciola hepatica',
        'trematodos - schistosoma sp.',  'trematodos schistosoma sp',
        'acantocefalos - macracanthorhynchus sp', 'acantocefalos macracanthorhynchus sp',
        'larvas de helmintos',
        'nematodos - ancylostoma sp.',   'nematodos ancylostoma sp',
        'nematodos - enterobius sp.',    'nematodos enterobius sp',
        'nematodos - strongyloides sp.', 'nematodos strongyloides sp',
        'cestodos - diphyllobothrium sp.','cestodos diphyllobothrium sp',
        'nematodos - trichuris sp.',     'nematodos trichuris sp',
        'nematodos - ascaris sp.',       'nematodos ascaris sp',
        'cestodos - taenia sp.',         'cestodos taenia sp',
        'cestodos - hymenolepis diminuta','cestodos hymenolepis diminuta',
        'cestodos - hymenolepis nana',   'cestodos hymenolepis nana',
        'numeracion de huevos y larvas de helmintos',
        'formas parasitarias quistes y ooquistes de protozoarios patogenos',
        'formas parasitarias',
    ],
    # Organismos de vida libre — agrupados en D.S., individuales en Pacific Control
    'organismos de vida libre': [
        'organismos de vida libre',
        'organismos de vida libre algas en todos sus estadios evolutivos',
        'organismos de vida libre protozoarios en todos sus estadios evolutivos',
        'organismos de vida libre copepodos en todos sus estadios evolutivos',
        'organismos de vida libre rotiferos en todos sus estadios evolutivos',
        'organismos de vida libre nematodos en todos sus estadios evolutivos',
    ],
}

# Valores LMP conocidos que openpyxl puede leer mal (ej. rango de pH → fecha)
HARDCODED_LMP = {
    'ph':                        ('6.5 a 8.5', 'Unidad de pH'),
    'potencial de hidrogeno':    ('6.5 a 8.5', 'Unidad de pH'),
    'potencial de hidrogeno ph': ('6.5 a 8.5', 'Unidad de pH'),
    'cloro total':               (5, 'mg/L'),    # mismo límite que cloro residual libre
    'temperatura':               ('-', '°C'),     # sin límite numérico en D.S.
}


def _lmp_desde_excel(path):
    """
    Lee LMP desde Excel.
    - Usa la hoja 'DS-031-2010' si existe (Excel DECRETOS de Pacific Control)
    - Si no existe, usa la hoja activa con detección automática de columnas
    - Agrega alias para diferencias de nombres entre el D.S. y los PDFs
    """
    wb   = openpyxl.load_workbook(path, read_only=True, data_only=True)
    lmp  = {}

    # Elegir hoja correcta
    hoja_ds031 = next((s for s in wb.sheetnames if '031' in s or 'DS-031' in s.upper()), None)
    ws = wb[hoja_ds031] if hoja_ds031 else wb.active

    rows = list(ws.iter_rows(values_only=True))

    # Detectar fila de encabezado buscando "Parámetros" o "Límite"
    header_idx = None
    idx_param = idx_lmp = idx_unid = None

    for i, row in enumerate(rows):
        row_norm = [normalizar(str(v or '')) for v in row]
        for j, h in enumerate(row_norm):
            if h in ('parametros', 'parametro', 'analisis', 'nombre'):
                idx_param = j
            elif 'limite' in h or 'permisible' in h or h == 'lmp':
                idx_lmp = j
            elif h in ('unidad', 'unidades', 'unit'):
                idx_unid = j
        if idx_param is not None and idx_lmp is not None:
            header_idx = i
            break

    # Fallback posicional si no se encontró encabezado
    if header_idx is None:
        for i, row in enumerate(rows):
            non_none = [v for v in row if v is not None]
            if len(non_none) >= 2:
                for j, v in enumerate(row):
                    if isinstance(v, str) and len(v.strip()) > 2 and idx_param is None:
                        idx_param = j
                    elif isinstance(v, (int, float)) and idx_lmp is None and j != idx_param:
                        idx_lmp = j
                if idx_param is not None:
                    header_idx = max(0, i - 1)
                    break

    if header_idx is None or idx_param is None:
        wb.close()
        return lmp

    if idx_lmp is None:
        idx_lmp = idx_param + 2 if idx_param + 2 <= 7 else idx_param + 1
    if idx_unid is None and abs(idx_lmp - idx_param) > 1:
        idx_unid = min(idx_param, idx_lmp) + 1

    # Leer filas de datos
    SKIP = {'parametros', 'parametro', 'analisis', 'limite maximo permisible',
            'lmp', 'unidad', 'unidades'}

    for row in rows[header_idx + 1:]:
        if not row or all(v is None for v in row):
            continue

        def get(idx):
            return row[idx] if idx is not None and idx < len(row) else None

        param_raw = get(idx_param)
        lmp_raw   = get(idx_lmp)
        unid_raw  = get(idx_unid)

        if param_raw is None:
            continue
        param = str(param_raw).strip()
        if not param:
            continue
        param_norm = normalizar(param)
        if param_norm in SKIP or len(param_norm) < 2 or 'decreto' in param_norm:
            continue

        unidad = str(unid_raw).strip() if unid_raw not in (None, '') else ''

        # Limpiar LMP: si openpyxl devolvió un datetime (ej. rango de pH), ignorar
        from datetime import datetime as _dt
        if isinstance(lmp_raw, _dt):
            lmp_val = None  # se reemplazará por HARDCODED_LMP si aplica
        else:
            lmp_val = lmp_raw

        # Convertir strings con comas decimales a float si es posible
        if isinstance(lmp_val, str):
            try:
                lmp_val = float(lmp_val.replace(',', '.'))
            except ValueError:
                pass  # mantener como string (ej. rangos "6.5 a 8.5")

        entry = (lmp_val, unidad)

        # Guardar con nombre del Decreto (normalizado)
        lmp[param_norm] = entry

        # Verificar si aplica algún alias y agregar entradas adicionales
        for dec_key, pdf_names in DECRETOS_ALIAS.items():
            if dec_key in param_norm or param_norm in dec_key:
                alias_entry = entry
                for pdf_name in pdf_names:
                    lmp[normalizar(pdf_name)] = alias_entry

    wb.close()

    # Aplicar valores hardcoded para parámetros problemáticos
    for key, val in HARDCODED_LMP.items():
        if key not in lmp or lmp[key][0] is None:
            lmp[key] = val

    return lmp


def _lmp_desde_pdf(path):
    """
    Extrae LMP de un PDF. Busca tablas con columnas Parámetro / Unidad / LMP
    o texto con formato similar.
    """
    lmp = {}
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            tablas = page.extract_tables()
            for tabla in tablas:
                if not tabla or len(tabla) < 2:
                    continue
                # Detectar fila de encabezado
                header = [str(c or '').strip().lower() for c in tabla[0]]
                idx_param = next((i for i, h in enumerate(header)
                                  if any(a in h for a in ('par', 'anal', 'nombre'))), None)
                idx_lmp   = next((i for i, h in enumerate(header)
                                  if any(a in h for a in ('lmp', 'lim', 'valor', 'vma'))), None)
                idx_unid  = next((i for i, h in enumerate(header)
                                  if any(a in h for a in ('uni',))), None)

                if idx_param is None or idx_lmp is None:
                    continue

                for row in tabla[1:]:
                    if not row or all(c is None for c in row):
                        continue
                    param = str(row[idx_param] or '').strip()
                    val   = row[idx_lmp]
                    unid  = str(row[idx_unid] or '').strip() if idx_unid is not None else ''
                    if param and val is not None:
                        lmp[normalizar(param)] = (val, unid)

    return lmp


def buscar_lmp(analisis, lmp_dict):
    """Primero coincidencia exacta (normalizada), luego parcial por substring."""
    key = normalizar(analisis)
    if key in lmp_dict:
        return lmp_dict[key]
    for k, v in lmp_dict.items():
        if len(key) > 3 and (key in k or k in key):
            return v
    return None


# ═══════════════════════════════════════════════════════════════════════════════
# COMPARACIÓN RESULTADO vs LMP
# ═══════════════════════════════════════════════════════════════════════════════

def _parsear_numero(s):
    s = str(s).strip().split()[0]
    menor = s.startswith('<')
    try:
        return float(s.lstrip('<').replace(',', '.')), menor
    except ValueError:
        return None, menor


def evaluar(resultado_str, lmp_val):
    if lmp_val is None or str(lmp_val).strip() == '':
        return ''
    res, res_menor = _parsear_numero(resultado_str)
    if res is None:
        return ''
    lmp_str = str(lmp_val).strip().replace(',', '.')

    # Rango "6.5 a 8.5" o "0.5 - 5"
    m = re.match(r'([\d.]+)\s+a\s+([\d.]+)', lmp_str) or \
        re.match(r'([\d.]+)\s*[-–]\s*([\d.]+)', lmp_str)
    if m:
        lmp_min, lmp_max = float(m.group(1)), float(m.group(2))
        if res_menor:
            return 'CONFORME' if res <= lmp_max else 'NO CONFORME'
        return 'CONFORME' if lmp_min <= res <= lmp_max else 'NO CONFORME'

    try:
        lmp_num = float(lmp_str)
    except ValueError:
        return ''

    if res_menor:
        # LMP=0 (Ausencia): "<X" = no detectado → CONFORME
        if lmp_num == 0:
            return 'CONFORME'
        return 'CONFORME' if res <= lmp_num else 'NO CONFORME'
    return 'CONFORME' if res <= lmp_num else 'NO CONFORME'


# Ruta al logo Pacific Control (para bloque de firmas SIN FONDO)
LOGO_PACIFIC_PATH = os.path.join(os.path.dirname(__file__), 'logo_pacific.png')


# ═══════════════════════════════════════════════════════════════════════════════
# GENERACIÓN DEL WORD (basado en el template)
# ═══════════════════════════════════════════════════════════════════════════════

def _limpiar_header(doc):
    """Vacía el header del documento (elimina logo y fondo de Pacific Control)."""
    for section in doc.sections:
        hdr_el = section.header._element
        for child in list(hdr_el):
            hdr_el.remove(child)
        hdr_el.append(OxmlElement('w:p'))  # párrafo vacío mínimo


def _footer_general_terms(doc):
    """Reemplaza el footer con el texto 'general terms' de Pacific Control."""
    from docx.shared import RGBColor as _RGB
    for section in doc.sections:
        ftr = section.footer
        ftr_el = ftr._element
        for child in list(ftr_el):
            ftr_el.remove(child)

        # Crear párrafo en el footer
        p_el = OxmlElement('w:p')
        ftr_el.append(p_el)

        def _add_run_xml(text, hex_color, size_half=14):
            r_el = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            fonts = OxmlElement('w:rFonts')
            fonts.set(qn('w:ascii'), 'Arial')
            fonts.set(qn('w:hAnsi'), 'Arial')
            rPr.append(fonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(size_half))
            rPr.append(sz)
            color = OxmlElement('w:color')
            color.set(qn('w:val'), hex_color)
            rPr.append(color)
            r_el.append(rPr)
            t_el = OxmlElement('w:t')
            t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t_el.text = text
            r_el.append(t_el)
            p_el.append(r_el)

        _add_run_xml(
            'Our general term and conditions are available in full '
            'www.pacificcontrol.us or at your request\n',
            '000000'
        )
        _add_run_xml(
            'Offices, Resident Inspectors, Joint Ventureships, and '
            'Representativs throughtout os the world',
            'FF6600'
        )

def _shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _quitar_bordes(tabla):
    for row in tabla.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcB = OxmlElement('w:tcBorders')
            for lado in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                el = OxmlElement(f'w:{lado}')
                el.set(qn('w:val'), 'nil')
                tcB.append(el)
            tcPr.append(tcB)


def _p(doc, texto='', bold=False, size=10, align=None, antes=0, despues=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(antes)
    p.paragraph_format.space_after  = Pt(despues)
    if align:
        p.alignment = align
    if texto:
        r = p.add_run(texto)
        r.bold      = bold
        r.font.size = Pt(size)
        r.font.name = 'Arial'
    return p


def _label_val(doc, label, valor, size=10, align=None, bold_valor=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(14)
    if align:
        p.alignment = align
    rl = p.add_run(f'{label}: ')
    rl.bold = True; rl.font.size = Pt(size); rl.font.name = 'Arial'
    rv = p.add_run(valor)
    rv.bold = bold_valor
    rv.font.size = Pt(size); rv.font.name = 'Arial'
    return p


def _celda(cell, texto, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color_hex=None):
    p = cell.paragraphs[0]
    p.clear()
    r = p.add_run(texto)
    r.bold      = bold
    r.font.size = Pt(size)
    r.font.name = 'Arial'
    p.alignment = align
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    if color_hex:
        r.font.color.rgb = RGBColor.from_string(color_hex)
    return r


def _insertar_firmas_template(doc):
    """
    Extrae el bloque VML de firmas del template y lo inserta en el documento generado.
    Preserva las imágenes de firma escaneadas y el layout exacto del original.
    """
    import copy
    from lxml import etree

    tpl = Document(TEMPLATE_PATH)
    tpl_body = tpl.element.body
    body     = doc.element.body
    sect_pr  = body.find(qn('w:sectPr'))

    # Buscar en el template los párrafos que contienen el dibujo VML de firmas
    # (párrafos con w:drawing o v:group — son los de las firmas escaneadas)
    firma_paras = []
    for child in tpl_body:
        tag = child.tag
        if tag not in (qn('w:p'), qn('w:tbl')):
            continue
        xml_str = etree.tostring(child, encoding='unicode')
        if 'v:group' in xml_str or ('w:drawing' in xml_str and 'v:' in xml_str):
            firma_paras.append(copy.deepcopy(child))

    if not firma_paras:
        # Fallback: usar texto simple si no se encontró el VML
        ft = doc.add_table(rows=4, cols=2)
        _quitar_bordes(ft)
        firmas = [
            ('JOSE ANDRES HUIMAN DIAZ', 'Analista', 'PACIFIC CONTROL SAC', 'CIP 193383'),
            ('MBI. JOEL CARLOS ELIAS PAREDES', 'Supervisor de Lab. Microbiología',
             'PACIFIC CONTROL SAC', 'CBP 13240'),
        ]
        for col_i, (nombre, cargo, empresa, codigo) in enumerate(firmas):
            for row_i, txt in enumerate(['____________________________', nombre,
                                          cargo, f'{empresa}  {codigo}']):
                cell = ft.rows[row_i].cells[col_i]
                p = cell.paragraphs[0]
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(txt)
                r.font.size = Pt(9); r.font.name = 'Arial'
                r.bold = (row_i == 1)
        return

    # Insertar los párrafos VML antes del sectPr
    idx = list(body).index(sect_pr) if sect_pr is not None else len(list(body))
    for para in firma_paras:
        if sect_pr is not None:
            body.insert(idx, para)
            idx += 1
        else:
            body.append(para)


def _insertar_firmas_sin_fondo(doc):
    """
    Bloque de firmas para formato SIN FONDO.
    Layout: [logo | línea + nombre/cargo] × 2 columnas (Jose / Joel).
    Sin rúbricas escaneadas — solo logo corporativo + datos.
    """
    import io as _io
    AZUL = RGBColor(0x00, 0x2B, 0x7F)   # azul Pacific Control

    FIRMAS = [
        ('JOSE ANDRES HUIMAN DIAZ',      'Analista',
         'PACIFIC CONTROL SAC',           'CIP 193383'),
        ('MBI. JOEL CARLOS ELIAS PAREDES','Supervisor de Lab. Microbiología',
         'PACIFIC CONTROL SAC',           'CBP 13240'),
    ]

    # Tabla exterior: 1 fila, 2 columnas (Jose | Joel)
    t_ext = doc.add_table(rows=1, cols=2)
    _quitar_bordes(t_ext)
    t_ext.alignment = WD_TABLE_ALIGNMENT.CENTER

    for col_i, (nombre, cargo, empresa, codigo) in enumerate(FIRMAS):
        cell_ext = t_ext.rows[0].cells[col_i]
        cell_ext.width = Cm(7.5)

        # Tabla interior: 1 fila, 2 columnas (logo | línea+texto)
        t_int = cell_ext.add_table(rows=1, cols=2)
        _quitar_bordes(t_int)

        # ── Celda izquierda: logo ──────────────────────────────────────────
        c_logo = t_int.rows[0].cells[0]
        c_logo.width = Cm(2.0)
        p_logo = c_logo.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(0)
        p_logo.paragraph_format.space_after  = Pt(0)
        if os.path.isfile(LOGO_PACIFIC_PATH):
            p_logo.add_run().add_picture(LOGO_PACIFIC_PATH, width=Cm(1.7))

        # ── Celda derecha: línea azul + nombre/cargo ───────────────────────
        c_txt = t_int.rows[0].cells[1]
        c_txt.width = Cm(5.5)

        # Línea (borde inferior del párrafo = línea de firma)
        p_linea = c_txt.paragraphs[0]
        p_linea.paragraph_format.space_before = Pt(0)
        p_linea.paragraph_format.space_after  = Pt(3)
        # Añadir borde inferior al párrafo para simular la línea de firma
        pPr = p_linea._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '002B7F')
        pBdr.append(bottom)
        pPr.append(pBdr)
        r = p_linea.add_run(' ')          # espacio para que la línea tenga altura
        r.font.size = Pt(12)

        def _txt_azul(cell, texto, bold=False, size=9):
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run(texto)
            r.bold = bold; r.font.size = Pt(size)
            r.font.name = 'Arial'
            r.font.color.rgb = AZUL

        _txt_azul(c_txt, nombre,   bold=True)
        _txt_azul(c_txt, cargo)
        _txt_azul(c_txt, empresa)
        _txt_azul(c_txt, codigo)


def _texto_conclusion(numeros, no_conformes):
    """Genera el texto por defecto de la conclusión según los resultados."""
    if not no_conformes:
        return (
            f'El resultado del análisis perteneciente al Informe de Ensayo N° {numeros}, '
            f'cumplen con los valores establecidos en el Decreto Supremo N° 031-2010-SA '
            f'Reglamento de la Calidad del Agua para Consumo Humano.'
        )
    params = ', '.join(no_conformes)
    return (
        f'Los resultados de los análisis del Informe de Ensayo N° {numeros}, '
        f'se encuentran dentro de los Límites Máximos Permisibles del Decreto Supremo '
        f'N° 031-2010-SA Reglamento de la Calidad del Agua para Consumo Humano, '
        f'excepto los parámetros: {params}.'
    )


def construir_datos(numero_ct, informes, lmp_dict):
    """
    Extrae el encabezado y evalúa cada parámetro contra los LMP, SIN generar el
    Word. Devuelve un dict listo para mostrarse en la pantalla de revisión y, una
    vez corregido por el usuario, pasarse a generar_word_desde_datos().
    """
    inf0    = informes[0]
    numeros = ' y '.join(i['numero'] for i in informes if i['numero'])

    encabezado = {
        'razon_social':   inf0['razon_social'],
        'direccion':      inf0['direccion'],
        'procedencia':    inf0['procedencia'],
        'cotizacion':     inf0['cotizacion'],
        'producto':       inf0['producto'],
        'punto_muestreo': inf0['punto_muestreo'],
        'presentacion':   inf0['presentacion'],
    }

    # Filas únicas (sin duplicados por nombre normalizado, respetando el orden)
    todos, vistos = [], set()
    for inf_i in informes:
        for fila in inf_i['resultados']:
            key = normalizar(fila[0])
            if key not in vistos:
                vistos.add(key)
                todos.append(fila)

    filas, sin_lmp = [], []
    for analisis, unidad, resultado in todos:
        lmp_info = buscar_lmp(analisis, lmp_dict)
        lmp_val  = lmp_info[0] if lmp_info else None
        ev       = evaluar(resultado, lmp_val)
        if lmp_info is None:
            sin_lmp.append(analisis)
        lmp_display = str(lmp_val).replace('.', ',') if lmp_val not in (None, '') else '-'
        res_display = resultado.split()[0] if resultado else resultado
        filas.append({
            'analisis':   analisis,
            'unidad':     unidad,
            'resultado':  res_display,
            'lmp':        lmp_display,
            'evaluacion': ev,
        })

    no_conf_inicial = [f['analisis'] for f in filas if f.get('evaluacion') == 'NO CONFORME']

    return {
        'numero':      numero_ct,
        'encabezado':  encabezado,
        'numeros':     numeros,
        'tiene_micro': any(i['tiene_micro'] for i in informes),
        'filas':       filas,
        'sin_lmp':     sin_lmp,
        'conclusion':  _texto_conclusion(numeros, no_conf_inicial),
    }


def generar_word_desde_datos(datos, output_path):
    """
    Genera el Comentario Técnico (.docx) a partir de un dict de datos ya evaluados
    (y posiblemente corregidos por el usuario en la pantalla de revisión).
    Preserva header (logo + fondo) y footer (número de página) del template.
    Retorna la lista de parámetros NO CONFORMES.
    """
    # ── Abrir template y limpiar cuerpo preservando sectPr ──────────────────
    doc = Document(TEMPLATE_PATH)
    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))

    # Eliminar todos los elementos excepto sectPr
    for el in list(body):
        if el is not sect_pr:
            body.remove(el)

    # ── Datos comunes ────────────────────────────────────────────────────────
    enc       = datos.get('encabezado', {})
    numeros   = datos.get('numeros', '')
    numero_ct = datos.get('numero', '')

    # ── Título ───────────────────────────────────────────────────────────────
    _p(doc, f'COMENTARIO TÉCNICO N° {numero_ct}',
       bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, antes=0, despues=8)

    # ── Encabezado ───────────────────────────────────────────────────────────
    _label_val(doc, 'RAZÓN SOCIAL',    enc.get('razon_social', ''))
    _label_val(doc, 'DIRECCIÓN LEGAL', enc.get('direccion', ''))
    _label_val(doc, 'PROCEDENCIA',     enc.get('procedencia', ''))
    _label_val(doc, 'COTIZACIÓN',      enc.get('cotizacion', ''))
    _label_val(doc, 'INFORME DE ENSAYO N°', numeros,
               align=WD_ALIGN_PARAGRAPH.CENTER, bold_valor=True)
    muestra = (f"{numeros} / {enc.get('producto', '')} / "
               f"{enc.get('punto_muestreo', '')} / {enc.get('presentacion', '')}")
    _label_val(doc, 'Muestra Id', muestra,
               align=WD_ALIGN_PARAGRAPH.CENTER, bold_valor=True)

    _p(doc)

    # ── Texto introductorio ───────────────────────────────────────────────────
    p_intro = doc.add_paragraph()
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_intro.paragraph_format.space_before = Pt(0)
    p_intro.paragraph_format.space_after  = Pt(6)
    r = p_intro.add_run(
        'A continuación, mediante cuadro comparativo (Cuadro N°1), se presenta la evaluación '
        'de los resultados de análisis obtenidos respecto al Decreto Supremo N° 031-2010-SA '
        'Reglamento de la Calidad del Agua para Consumo Humano'
    )
    r.font.size = Pt(10); r.font.name = 'Arial'

    _p(doc)

    # ── Cuadro N°1 ────────────────────────────────────────────────────────────
    _p(doc, 'Cuadro N°1', bold=True, size=10,
       align=WD_ALIGN_PARAGRAPH.CENTER, antes=0, despues=2)
    _p(doc,
       'Resultados de Laboratorio - Son comparados con los valores respectivos al Decreto '
       'Supremo N° 031-2010-SA Reglamento de la Calidad del Agua para Consumo Humano',
       size=9, align=WD_ALIGN_PARAGRAPH.CENTER, antes=0, despues=4)

    # ── Tabla de resultados ───────────────────────────────────────────────────
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    HDRS    = ['Análisis', 'Unidad', 'Resultados', 'D.S. N° 031-2010-SA', 'Evaluación']
    COL_CM  = [5.2, 2.2, 2.2, 2.9, 2.0]
    C_HDR   = '002060'  # azul marino del template original

    for cell, h in zip(table.rows[0].cells, HDRS):
        _celda(cell, h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color_hex='FFFFFF')
        _shading(cell, C_HDR)

    no_conformes = []

    for fila in datos.get('filas', []):
        analisis  = (fila.get('analisis')   or '').strip()
        unidad    = (fila.get('unidad')     or '').strip()
        resultado = (fila.get('resultado')  or '').strip()
        lmp_disp  = (fila.get('lmp')         or '').strip() or '-'
        ev        = (fila.get('evaluacion')  or '').strip()

        if not analisis and not resultado:
            continue  # ignorar filas vacías que el usuario no llegó a llenar

        if ev == 'NO CONFORME':
            no_conformes.append(analisis)

        row_cells = table.add_row().cells
        vals   = [analisis, unidad, resultado, lmp_disp, ev]
        aligns = [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 4

        for i, (cell, val, aln) in enumerate(zip(row_cells, vals, aligns)):
            # NO CONFORME → resultado (col 2) y evaluación (col 4) en azul negrita
            if ev == 'NO CONFORME' and i in (2, 4):
                _celda(cell, val, bold=True, align=aln, color_hex='00B0F0')
            else:
                _celda(cell, val, align=aln)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(COL_CM[i])

    _p(doc)

    # ── Conclusión (texto aprobado por el usuario en la pantalla de revisión) ──
    p_cl = doc.add_paragraph()
    p_cl.paragraph_format.space_before = Pt(0)
    p_cl.paragraph_format.space_after  = Pt(2)
    p_cl.add_run('Conclusión:').bold = True

    p_cl2 = doc.add_paragraph()
    p_cl2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_cl2.paragraph_format.space_before = Pt(0)
    p_cl2.paragraph_format.space_after  = Pt(4)

    concl = datos.get('conclusion') or _texto_conclusion(numeros, no_conformes)
    r = p_cl2.add_run(concl)
    r.font.size = Pt(10); r.font.name = 'Arial'

    # ── Referencias ───────────────────────────────────────────────────────────
    _p(doc)
    p_ref = doc.add_paragraph()
    p_ref.paragraph_format.space_before = Pt(0)
    p_ref.paragraph_format.space_after  = Pt(2)
    p_ref.add_run('Referencias:').bold = True

    _p(doc,
       'Decreto Supremo N° 031-2010-SA Reglamento de la Calidad del Agua para Consumo Humano',
       size=10, antes=0, despues=4)

    # ── Fecha ─────────────────────────────────────────────────────────────────
    _p(doc)
    h = datetime.today()
    _p(doc, f'Lima, {h.day} de {MESES[h.month]} del {h.year}',
       size=10, align=WD_ALIGN_PARAGRAPH.RIGHT, antes=0, despues=0)

    _p(doc); _p(doc)

    # ── Firmas: copiar bloque VML exacto del template ─────────────────────────
    # Las imágenes de firma son flotantes (wp:anchor). En el template original
    # hay 11 párrafos vacíos después del VML para que "FIN DEL DOCUMENTO" no
    # quede tapado por las imágenes flotantes.
    _insertar_firmas_template(doc)

    for _ in range(11):
        _p(doc)
    _p(doc, 'FIN DEL DOCUMENTO', bold=True, size=10,
       align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(output_path)
    return no_conformes


def generar_word_sin_fondo(datos, output_path):
    """
    Genera el Comentario Técnico en formato SIN FONDO:
    fondo blanco limpio, sin header de Pacific Control, footer 'general terms'.
    Usa el template CON FONDO como base para que las imágenes de firma
    (rId11, rId13) ya estén en el documento y el VML funcione correctamente.
    Retorna la lista de parámetros NO CONFORMES.
    """
    # ── Base: template CON FONDO (preserva relaciones de imágenes de firma) ──
    doc = Document(TEMPLATE_PATH)

    # Vaciar header y cambiar footer ANTES de limpiar el cuerpo
    _limpiar_header(doc)
    _footer_general_terms(doc)

    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))

    for el in list(body):
        if el is not sect_pr:
            body.remove(el)

    # ── Todos los datos ──────────────────────────────────────────────────────
    enc       = datos.get('encabezado', {})
    numeros   = datos.get('numeros', '')
    numero_ct = datos.get('numero', '')

    # ── Título ───────────────────────────────────────────────────────────────
    _p(doc, f'COMENTARIO TÉCNICO N° {numero_ct}',
       bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, antes=0, despues=8)

    # ── Encabezado ───────────────────────────────────────────────────────────
    _label_val(doc, 'RAZÓN SOCIAL',         enc.get('razon_social', ''))
    _label_val(doc, 'DIRECCIÓN LEGAL',      enc.get('direccion', ''))
    _label_val(doc, 'PROCEDENCIA',          enc.get('procedencia', ''))
    _label_val(doc, 'COTIZACIÓN',           enc.get('cotizacion', ''))
    _label_val(doc, 'INFORME DE ENSAYO N°', numeros,
               align=WD_ALIGN_PARAGRAPH.CENTER, bold_valor=True)
    muestra = (f"{numeros} / {enc.get('producto', '')} / "
               f"{enc.get('punto_muestreo', '')} / {enc.get('presentacion', '')}")
    _label_val(doc, 'Muestra Id', muestra,
               align=WD_ALIGN_PARAGRAPH.CENTER, bold_valor=True)

    _p(doc)

    # ── Texto introductorio ───────────────────────────────────────────────────
    p_intro = doc.add_paragraph()
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_intro.paragraph_format.space_before = Pt(0)
    p_intro.paragraph_format.space_after  = Pt(6)
    r = p_intro.add_run(
        'A continuación, mediante cuadro comparativo (Cuadro N°1), se presenta la evaluación '
        'de los resultados de análisis obtenidos respecto al Decreto Supremo N° 031-2010-SA '
        'Reglamento de la Calidad del Agua para Consumo Humano'
    )
    r.font.size = Pt(10); r.font.name = 'Arial'

    _p(doc)

    # ── Cuadro N°1 ────────────────────────────────────────────────────────────
    _p(doc, 'Cuadro N°1', bold=True, size=10,
       align=WD_ALIGN_PARAGRAPH.CENTER, antes=0, despues=2)
    _p(doc,
       'Resultados de Laboratorio - Son comparados con los valores respectivos al Decreto '
       'Supremo N° 031-2010-SA Reglamento de la Calidad del Agua para Consumo Humano',
       bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, antes=0, despues=4)

    # ── Tabla de resultados ───────────────────────────────────────────────────
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    HDRS   = ['Análisis', 'Unidad', 'Resultados', 'D.S. N° 031-2010-SA', 'Evaluación']
    COL_CM = [5.2, 2.2, 2.2, 2.9, 2.0]
    C_HDR  = '002060'

    for cell, h in zip(table.rows[0].cells, HDRS):
        _celda(cell, h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color_hex='FFFFFF')
        _shading(cell, C_HDR)

    no_conformes = []

    for fila in datos.get('filas', []):
        analisis  = (fila.get('analisis')   or '').strip()
        unidad    = (fila.get('unidad')     or '').strip()
        resultado = (fila.get('resultado')  or '').strip()
        lmp_disp  = (fila.get('lmp')        or '').strip() or '-'
        ev        = (fila.get('evaluacion') or '').strip()

        if not analisis and not resultado:
            continue

        if ev == 'NO CONFORME':
            no_conformes.append(analisis)

        row_cells = table.add_row().cells
        vals   = [analisis, unidad, resultado, lmp_disp, ev]
        aligns = [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 4

        for i, (cell, val, aln) in enumerate(zip(row_cells, vals, aligns)):
            # NO CONFORME → resultado (col 2) y evaluación (col 4) en azul negrita
            if ev == 'NO CONFORME' and i in (2, 4):
                _celda(cell, val, bold=True, align=aln, color_hex='00B0F0')
            else:
                _celda(cell, val, align=aln)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(COL_CM[i])

    _p(doc)

    # ── Conclusión ────────────────────────────────────────────────────────────
    p_cl = doc.add_paragraph()
    p_cl.paragraph_format.space_before = Pt(0)
    p_cl.paragraph_format.space_after  = Pt(2)
    p_cl.add_run('Conclusión:').bold = True

    p_cl2 = doc.add_paragraph()
    p_cl2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_cl2.paragraph_format.space_before = Pt(0)
    p_cl2.paragraph_format.space_after  = Pt(4)
    concl = datos.get('conclusion') or _texto_conclusion(numeros, no_conformes)
    r = p_cl2.add_run(concl)
    r.font.size = Pt(10); r.font.name = 'Arial'

    # ── Referencias ───────────────────────────────────────────────────────────
    _p(doc)
    p_ref = doc.add_paragraph()
    p_ref.paragraph_format.space_before = Pt(0)
    p_ref.paragraph_format.space_after  = Pt(2)
    p_ref.add_run('Referencias:').bold = True

    _p(doc,
       'Decreto Supremo N° 031-2010-SA Reglamento de la Calidad del Agua para Consumo Humano',
       size=10, antes=0, despues=4)

    # ── Fecha ─────────────────────────────────────────────────────────────────
    _p(doc)
    h = datetime.today()
    _p(doc, f'Lima, {h.day} de {MESES[h.month]} del {h.year}',
       size=10, align=WD_ALIGN_PARAGRAPH.RIGHT, antes=0, despues=0)

    _p(doc); _p(doc)

    # ── Firmas SIN FONDO: logo + línea + texto (sin rúbricas escaneadas) ──────
    _insertar_firmas_sin_fondo(doc)

    _p(doc)
    _p(doc, 'FIN DEL DOCUMENTO', bold=True, size=10,
       align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(output_path)
    return no_conformes


def generar_word(numero_ct, informes, lmp_dict, output_path):
    """
    Compatibilidad: construye los datos y genera el Word en un solo paso.
    Retorna (lista_no_conformes, lista_sin_lmp).
    """
    datos = construir_datos(numero_ct, informes, lmp_dict)
    no_conformes = generar_word_desde_datos(datos, output_path)
    return no_conformes, datos['sin_lmp']


# ═══════════════════════════════════════════════════════════════════════════════
# GENERACIÓN DE PDF (fpdf2)
# ═══════════════════════════════════════════════════════════════════════════════

def generar_pdf(datos, output_path, con_fondo=True):
    """
    Genera el Comentario Técnico en formato PDF usando fpdf2.
    con_fondo=True  → formato CON FONDO (header/background Pacific Control)
    con_fondo=False → formato SIN FONDO (blanco limpio, footer general terms)
    """
    from fpdf import FPDF
    import zipfile, io as _io

    # ── Colores ──────────────────────────────────────────────────────────────
    C_AZUL    = (0, 32, 96)        # 002060 — header tabla
    C_CELESTE = (0, 176, 240)      # 00B0F0 — NO CONFORME
    C_BLANCO  = (255, 255, 255)
    C_NEGRO   = (0, 0, 0)
    C_NARANJA = (255, 102, 0)      # footer sin fondo
    C_AZUL_F  = (0, 43, 127)       # texto firmas

    # ── Clase PDF con footer ─────────────────────────────────────────────────
    class _PDF(FPDF):
        def __init__(self, sf):
            super().__init__()
            self._sf = sf           # True = sin fondo

        def header(self):
            pass                    # el header se inserta manualmente en add_page

        def footer(self):
            fw = self.w - self.l_margin - self.r_margin
            self.set_y(-18)
            if self._sf:
                self.set_font('Helvetica', '', 7)
                self.set_text_color(*C_NEGRO)
                self.cell(fw, 3.5,
                    'Our general term and conditions are available in full '
                    'www.pacificcontrol.us or at your request',
                    new_x='LMARGIN', new_y='NEXT')
                self.set_font('Helvetica', '', 7)
                self.set_text_color(*C_NARANJA)
                self.cell(fw, 3.5,
                    'Offices, Resident Inspectors, Joint Ventureships, and '
                    'Representativs throughtout os the world',
                    new_x='LMARGIN', new_y='NEXT')
            else:
                self.set_font('Helvetica', '', 8)
                self.set_text_color(100, 100, 100)
                self.cell(fw, 5, f'Página {self.page_no()}', align='C')

    # ── Crear PDF ─────────────────────────────────────────────────────────────
    pdf = _PDF(sf=not con_fondo)
    pdf.set_margins(30, 28 if con_fondo else 25, 30)
    pdf.set_auto_page_break(auto=True, margin=22)

    # Registrar fuentes con soporte UTF-8
    _fname = 'Helvetica'   # fallback
    if os.path.isfile(_FONT_REGULAR):
        pdf.add_font('ct_regular', '', _FONT_REGULAR)
        _fname = 'ct_regular'
    if os.path.isfile(_FONT_BOLD):
        pdf.add_font('ct_bold', '', _FONT_BOLD)
    _fbold = 'ct_bold' if os.path.isfile(_FONT_BOLD) else _fname

    def _font(bold=False, size=10):
        name = _fbold if bold else _fname
        pdf.set_font(name, '', size)

    pdf.add_page()
    W = pdf.w - pdf.l_margin - pdf.r_margin   # ancho útil ≈ 150 mm

    # ── Background / header (CON FONDO) ───────────────────────────────────────
    if con_fondo:
        try:
            with zipfile.ZipFile(TEMPLATE_PATH) as z:
                bg = z.read('word/media/image5.jpg')
            # Fondo completo de la página
            pdf.image(_io.BytesIO(bg), x=0, y=0, w=pdf.w, h=pdf.h)
            # Reiniciar posición después de la imagen de fondo
            pdf.set_xy(pdf.l_margin, 28)
        except Exception:
            pass

    enc       = datos.get('encabezado', {})
    numeros   = datos.get('numeros', '')
    numero_ct = str(datos.get('numero', '')).strip()

    # ── Titulo ───────────────────────────────────────────────────────────────
    _font(bold=True, size=12)
    pdf.set_text_color(*C_NEGRO)
    pdf.cell(W, 8, f'COMENTARIO TÉCNICO N° {numero_ct}', align='C', new_x='LMARGIN', new_y='NEXT')
    pdf.ln(3)

    # ── Encabezado ────────────────────────────────────────────────────────────
    def _lv(label, valor, center=False):
        _font(bold=True, size=10)
        lw = pdf.get_string_width(f'{label}: ') + 2
        if center:
            pdf.cell(lw, 5.5, f'{label}: ', new_x='RIGHT', new_y='TOP')
            pdf.multi_cell(W - lw, 5.5, valor, align='C', new_x='LMARGIN', new_y='NEXT')
        else:
            pdf.cell(lw, 5.5, f'{label}: ', new_x='RIGHT', new_y='TOP')
            _font(bold=False, size=10)
            pdf.multi_cell(W - lw, 5.5, valor, new_x='LMARGIN', new_y='NEXT')

    _lv('RAZÓN SOCIAL',    enc.get('razon_social', ''))
    _lv('DIRECCIÓN LEGAL', enc.get('direccion', ''))
    _lv('PROCEDENCIA',     enc.get('procedencia', ''))
    _lv('COTIZACIÓN',      enc.get('cotizacion', ''))
    pdf.ln(2)

    _font(bold=True, size=10)
    pdf.cell(W, 5.5, f'INFORME DE ENSAYO N° {numeros}', align='C', new_x='LMARGIN', new_y='NEXT')
    muestra = (f'{numeros} / {enc.get("producto", "")} / '
               f'{enc.get("punto_muestreo", "")} / {enc.get("presentacion", "")}')
    pdf.multi_cell(W, 5.5, f'Muestra Id: {muestra}', align='C', new_x='LMARGIN', new_y='NEXT')
    pdf.ln(3)

    # ── Texto introductorio ───────────────────────────────────────────────────
    _font(bold=False, size=10)
    pdf.set_text_color(*C_NEGRO)
    pdf.multi_cell(W, 5,
        'A continuación, mediante cuadro comparativo (Cuadro N°1), se presenta la evaluación '
        'de los resultados de análisis obtenidos respecto al Decreto Supremo N° 031-2010-SA '
        'Reglamento de la Calidad del Agua para Consumo Humano', align='J',
        new_x='LMARGIN', new_y='NEXT')
    pdf.ln(3)

    # ── Cuadro N°1 ────────────────────────────────────────────────────────────
    _font(bold=True, size=10)
    pdf.cell(W, 5, 'Cuadro N°1', align='C', new_x='LMARGIN', new_y='NEXT')
    _font(bold=not con_fondo, size=9)
    pdf.multi_cell(W, 4.5,
        'Resultados de Laboratorio - Son comparados con los valores respectivos al Decreto '
        'Supremo N° 031-2010-SA Reglamento de la Calidad del Agua para Consumo Humano',
        align='C', new_x='LMARGIN', new_y='NEXT')
    pdf.ln(2)

    # ── Tabla de resultados ───────────────────────────────────────────────────
    COL_W  = [52, 22, 22, 29, 20]     # mm (total ≈ 145 mm)
    HDRS   = ['Análisis', 'Unidad', 'Resultados', 'D.S. N° 031-2010-SA', 'Evaluación']
    LINE_H = 5.0

    def _dibujar_cabecera():
        pdf.set_fill_color(*C_AZUL)
        pdf.set_text_color(*C_BLANCO)
        _font(bold=True, size=8)
        for h, w in zip(HDRS, COL_W):
            pdf.cell(w, 7, h, border=1, align='C', fill=True)
        pdf.ln()

    _dibujar_cabecera()

    for fila in datos.get('filas', []):
        analisis  = (fila.get('analisis')   or '').strip()
        unidad    = (fila.get('unidad')     or '').strip()
        resultado = (fila.get('resultado')  or '').strip()
        lmp_disp  = (fila.get('lmp')        or '').strip() or '-'
        ev        = (fila.get('evaluacion') or '').strip()

        if not analisis and not resultado:
            continue

        es_nc = (ev == 'NO CONFORME')

        # ── Estimar altura de la fila (en base al ancho de texto) ────────────
        _font(bold=False, size=8)
        txt_w = pdf.get_string_width(analisis)
        n_lin = max(1, int(txt_w / (COL_W[0] - 2)) + 1)
        row_h = n_lin * LINE_H

        # Salto de página con re-cabecera si es necesario
        if pdf.get_y() + row_h > pdf.page_break_trigger:
            pdf.add_page()
            if con_fondo:
                try:
                    with zipfile.ZipFile(TEMPLATE_PATH) as z:
                        bg = z.read('word/media/image5.jpg')
                    pdf.image(_io.BytesIO(bg), x=0, y=0, w=pdf.w, h=pdf.h)
                    pdf.set_xy(pdf.l_margin, 28)
                except Exception:
                    pass
            _dibujar_cabecera()

        y0 = pdf.get_y()
        x0 = pdf.l_margin

        # Col 0: Análisis (multi-línea)
        pdf.set_xy(x0, y0)
        _font(bold=False, size=8)
        pdf.set_text_color(*C_NEGRO)
        pdf.multi_cell(COL_W[0], LINE_H, analisis, border='LTB', align='L')
        row_h = pdf.get_y() - y0

        # Cols 1-4
        for ci, val, celeste in [(1, unidad, False), (2, resultado, es_nc),
                                  (3, lmp_disp, False), (4, ev, es_nc)]:
            pdf.set_xy(x0 + sum(COL_W[:ci]), y0)
            _font(bold=celeste, size=8)
            pdf.set_text_color(*(C_CELESTE if celeste else C_NEGRO))
            pdf.cell(COL_W[ci], row_h, val, border=1, align='C')

        pdf.set_y(y0 + row_h)

    pdf.ln(4)

    # ── Conclusión ────────────────────────────────────────────────────────────
    _font(bold=True, size=10);  pdf.set_text_color(*C_NEGRO)
    pdf.cell(W, 5, 'Conclusión:', new_x='LMARGIN', new_y='NEXT')
    _font(bold=False, size=10)
    concl = datos.get('conclusion') or _texto_conclusion(numeros, [])
    pdf.multi_cell(W, 5, concl, align='J', new_x='LMARGIN', new_y='NEXT')
    pdf.ln(3)

    # ── Referencias ───────────────────────────────────────────────────────────
    _font(bold=True, size=10)
    pdf.cell(W, 5, 'Referencias:', new_x='LMARGIN', new_y='NEXT')
    _font(bold=False, size=10)
    pdf.multi_cell(W, 5,
        'Decreto Supremo N° 031-2010-SA Reglamento de la Calidad del Agua para Consumo Humano',
        new_x='LMARGIN', new_y='NEXT')
    pdf.ln(3)

    # ── Fecha ─────────────────────────────────────────────────────────────────
    h = datetime.today()
    _font(bold=False, size=10)
    pdf.cell(W, 5, f'Lima, {h.day} de {MESES[h.month]} del {h.year}',
             align='R', new_x='LMARGIN', new_y='NEXT')
    pdf.ln(8)

    # ── Firmas ────────────────────────────────────────────────────────────────
    FIRMAS = [
        ('JOSE ANDRES HUIMAN DIAZ',       'Analista',
         'PACIFIC CONTROL SAC',            'CIP 193383'),
        ('MBI. JOEL CARLOS ELIAS PAREDES', 'Supervisor de Lab. Microbiología',
         'PACIFIC CONTROL SAC',            'CBP 13240'),
    ]
    bloque_w = W / 2 - 5
    y_firma  = pdf.get_y()

    if con_fondo:
        # CON FONDO: imágenes completas del template (sello + firma escaneada + texto)
        try:
            with zipfile.ZipFile(TEMPLATE_PATH) as z:
                img_jose = z.read('word/media/image1.png')
                img_joel = z.read('word/media/image3.png')
            x_jose = pdf.l_margin
            x_joel = pdf.l_margin + bloque_w + 10
            pdf.image(_io.BytesIO(img_jose), x=x_jose, y=y_firma, w=bloque_w)
            pdf.image(_io.BytesIO(img_joel), x=x_joel, y=y_firma, w=bloque_w)
            # Altura proporcional: max(0.528, 0.594) × bloque_w + margen
            img_h = bloque_w * 0.60
            pdf.set_y(y_firma + img_h + 5)
        except Exception:
            pdf.set_y(y_firma + 30)
    else:
        # SIN FONDO: solo sello + línea + texto (sin firma escaneada)
        logo_w = 18
        for idx, (nombre, cargo, empresa, codigo) in enumerate(FIRMAS):
            bx  = pdf.l_margin + idx * (bloque_w + 10)
            lx1 = bx + logo_w + 3
            lx2 = bx + bloque_w

            if os.path.isfile(LOGO_PACIFIC_PATH):
                pdf.image(LOGO_PACIFIC_PATH, x=bx, y=y_firma, h=logo_w)

            ly = y_firma + logo_w - 1
            pdf.set_draw_color(*C_AZUL_F)
            pdf.line(lx1, ly, lx2, ly)

            pdf.set_xy(lx1, ly + 1.5)
            _font(bold=True, size=8);  pdf.set_text_color(*C_AZUL_F)
            pdf.cell(lx2 - lx1, 4, nombre, align='L', new_x='LMARGIN', new_y='NEXT')
            for txt in [cargo, empresa, codigo]:
                _font(bold=False, size=8)
                pdf.set_xy(lx1, pdf.get_y())
                pdf.cell(lx2 - lx1, 4, txt, align='L', new_x='LMARGIN', new_y='NEXT')

            if idx == 0:
                pdf.set_y(y_firma)

        pdf.set_y(y_firma + logo_w + 20)

    # ── FIN DEL DOCUMENTO ─────────────────────────────────────────────────────
    _font(bold=True, size=10);  pdf.set_text_color(*C_NEGRO)
    pdf.cell(W, 5, 'FIN DEL DOCUMENTO', align='C', new_x='LMARGIN', new_y='NEXT')

    # output() devuelve bytearray; guardarlo según el tipo de output_path
    pdf_bytes = bytes(pdf.output())
    if hasattr(output_path, 'write'):
        output_path.write(pdf_bytes)
    else:
        with open(output_path, 'wb') as f:
            f.write(pdf_bytes)
